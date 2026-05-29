"""Competitive Analysis deliverable → DOCX + XLSX export.

Both formats open natively in Google Docs / Google Sheets via upload, and in
Word / Pages / Excel. Designed to match the in-app deliverable structure but
with a print/document-friendly professional color scheme.

Palette (intentional consultancy aesthetic):
  - Navy headings        #0F172A (rgb 15, 23, 42)
  - Charcoal body        #1E293B (rgb 30, 41, 59)
  - Slate captions       #475569 (rgb 71, 85, 105)
  - Amber accents        #B45309 (rgb 180, 83, 9)
  - Emerald positives    #047857 (rgb 4, 120, 87)
  - Background highlight #F8FAFC (rgb 248, 250, 252)
"""
from __future__ import annotations

from io import BytesIO
from typing import Any, Dict, List, Optional

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

NAVY = RGBColor(0x0F, 0x17, 0x2A)
CHARCOAL = RGBColor(0x1E, 0x29, 0x3B)
SLATE = RGBColor(0x47, 0x55, 0x69)
MUTED = RGBColor(0x94, 0xA3, 0xB8)
AMBER = RGBColor(0xB4, 0x53, 0x09)
EMERALD = RGBColor(0x04, 0x78, 0x57)
LIGHT_BG = "F8FAFC"
DIVIDER = "E2E8F0"
AMBER_BG = "FEF3C7"
EMERALD_BG = "D1FAE5"


# ----------------------------- helpers ---------------------------------------

def _fmt(n: Any) -> str:
    if n is None or n == "":
        return "—"
    if isinstance(n, (int, float)):
        if n >= 1_000_000:
            return f"{n / 1_000_000:.1f}M"
        if n >= 1_000:
            return f"{n / 1_000:.0f}k" if n >= 10_000 else f"{n / 1_000:.1f}k"
        return f"{int(n):,}" if isinstance(n, float) and n == int(n) else f"{n:,}"
    return str(n)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_cell_border(cell, color: str = DIVIDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def _add_heading(doc: Document, text: str, level: int = 1, color: RGBColor = NAVY) -> None:
    sizes = {1: 18, 2: 13, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(sizes.get(level, 11))
    run.font.color.rgb = color


def _add_kicker(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = AMBER
    run.bold = True
    # letter-spacing not natively supported; emulate with extra spaces? skip


def _add_body(doc: Document, text: str, color: RGBColor = CHARCOAL, italic: bool = False, size: int = 10) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if italic:
        run.italic = True


def _add_bullet(doc: Document, text: str, color: RGBColor = CHARCOAL) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = color


def _add_label_value(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = SLATE
    r2 = p.add_run(value)
    r2.font.size = Pt(10)
    r2.font.color.rgb = CHARCOAL


# ----------------------------- DOCX ------------------------------------------

def build_competitive_docx(content: Dict[str, Any], client_name: str = "") -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Default font for whole doc
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = CHARCOAL

    # ---------- Cover ----------
    _add_kicker(doc, "Competitive Analysis · Deliverable")
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run(content.get("title") or f"Competitive Analysis · {client_name}")
    title_run.bold = True
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = NAVY
    if content.get("subtitle"):
        _add_body(doc, content["subtitle"], color=SLATE, size=11)
    _add_body(doc, f"Prepared for {content.get('prepared_for') or client_name}", color=MUTED, italic=True, size=9)

    # ---------- Executive Summary ----------
    if content.get("executive_summary"):
        _add_heading(doc, "Executive Summary", 1)
        _add_body(doc, content["executive_summary"], size=11)

    # ---------- Current Position ----------
    cp = content.get("current_position") or {}
    if cp:
        _add_heading(doc, "Where You Stand Today", 1)
        table = doc.add_table(rows=1, cols=5)
        table.alignment = WD_ALIGN_PARAGRAPH.LEFT
        headers = ["Authority", "Backlinks", "Ref. domains", "Organic KWs", "Organic traffic"]
        values = [
            cp.get("authority_score"),
            cp.get("backlinks"),
            cp.get("referring_domains"),
            cp.get("organic_keywords"),
            cp.get("organic_traffic"),
        ]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            _set_cell_bg(cell, LIGHT_BG)
            _set_cell_border(cell)
            p = cell.paragraphs[0]
            r = p.add_run(h.upper())
            r.font.size = Pt(8)
            r.font.color.rgb = SLATE
            r.bold = True
        value_row = table.add_row()
        for i, v in enumerate(values):
            cell = value_row.cells[i]
            cell.text = ""
            _set_cell_border(cell)
            p = cell.paragraphs[0]
            r = p.add_run(_fmt(v))
            r.bold = True
            r.font.size = Pt(14)
            r.font.color.rgb = NAVY

        if cp.get("narrative"):
            doc.add_paragraph()
            _add_body(doc, cp["narrative"], size=11)

    # ---------- Competitor Landscape ----------
    landscape = content.get("competitor_landscape") or []
    if landscape:
        _add_heading(doc, "Competitor Landscape", 1)
        for comp in landscape:
            _add_heading(doc, f"{comp.get('name','?')}  ·  {comp.get('domain','?')}", 2, color=NAVY)
            if comp.get("tier"):
                _add_body(doc, comp["tier"].replace("_", " ").title(), color=AMBER, size=9)
            _add_label_value(doc, "Authority Score", str(comp.get("authority_score") or "—"))
            _add_label_value(doc, "Organic traffic (monthly)", _fmt(comp.get("organic_traffic")))
            if comp.get("key_strengths"):
                _add_body(doc, "Strengths:", color=EMERALD, size=10)
                for s in comp["key_strengths"]:
                    _add_bullet(doc, s)
            if comp.get("key_weaknesses"):
                _add_body(doc, "Weaknesses:", color=AMBER, size=10)
                for s in comp["key_weaknesses"]:
                    _add_bullet(doc, s)
            if comp.get("how_to_beat_them"):
                _add_label_value(doc, "How to beat them", comp["how_to_beat_them"])

    # ---------- Gap Analysis ----------
    gaps = content.get("gap_analysis") or {}
    if gaps:
        _add_heading(doc, "Gap Analysis", 1)
        for key, g in gaps.items():
            if not g:
                continue
            _add_heading(doc, g.get("label") or key.replace("_", " ").title(), 3, color=AMBER)
            you = g.get("your_value") if g.get("your_value") is not None else g.get("count")
            them = g.get("their_value") if g.get("their_value") is not None else g.get("count")
            _add_label_value(doc, "You", _fmt(you))
            if them is not None and them != you:
                _add_label_value(doc, f"{g.get('leader') or 'Leader'}", _fmt(them))
            if g.get("implication"):
                _add_body(doc, g["implication"])

    # ---------- Top Opportunities ----------
    opps = content.get("top_opportunities") or []
    if opps:
        _add_heading(doc, "Top Opportunities", 1)
        _add_body(doc, "Ranked by impact vs effort. Start at the top.", color=SLATE, italic=True, size=9)
        for opp in opps:
            rank = opp.get("rank") or "•"
            _add_heading(doc, f"#{rank}  {opp.get('title','?')}", 2)
            if opp.get("primary_keyword"):
                bits = [f"Primary: {opp['primary_keyword']}"]
                if opp.get("search_volume") is not None:
                    bits.append(f"Volume {_fmt(opp['search_volume'])}")
                if opp.get("competitor_position") is not None:
                    bits.append(f"They rank #{opp['competitor_position']}")
                _add_body(doc, "  ·  ".join(bits), color=EMERALD, size=9)
            if opp.get("supporting_keywords"):
                _add_body(doc, "Supporting: " + " · ".join(opp["supporting_keywords"][:6]), color=SLATE, size=9)
            if opp.get("why_winnable"):
                _add_body(doc, opp["why_winnable"])
            _add_label_value(doc, "Format", opp.get("recommended_format") or "—")
            _add_label_value(doc, "Effort / Impact", f"{opp.get('effort','—')} effort  ·  {opp.get('expected_impact','—')} impact")
            if opp.get("competitor_url_to_outrank"):
                _add_label_value(doc, "URL to outrank", opp["competitor_url_to_outrank"])
            if opp.get("first_step"):
                _add_label_value(doc, "First step", opp["first_step"])

    # ---------- Content Strategy ----------
    cs = content.get("content_strategy") or {}
    if cs:
        _add_heading(doc, "Content Strategy", 1)
        if cs.get("positioning"):
            _add_body(doc, cs["positioning"], size=11)
        if cs.get("pillars"):
            _add_heading(doc, "Content pillars", 3)
            for p in cs["pillars"]:
                _add_body(doc, p.get("name", "?"), size=10)
                if p.get("rationale"):
                    _add_body(doc, p["rationale"], color=SLATE, size=9)
                if p.get("example_targets"):
                    _add_body(doc, "Targets: " + " · ".join(p["example_targets"]), color=EMERALD, size=9, italic=True)
        if cs.get("format_priorities"):
            _add_heading(doc, "Format priorities", 3)
            for fp in cs["format_priorities"]:
                _add_bullet(doc, fp)
        if cs.get("voice_and_eeat"):
            _add_label_value(doc, "Voice & E-E-A-T", cs["voice_and_eeat"])

    # ---------- Link-building Strategy ----------
    lb = content.get("link_building_strategy") or {}
    if lb:
        _add_heading(doc, "Link-Building Strategy", 1)
        if lb.get("current_gap_summary"):
            _add_body(doc, lb["current_gap_summary"])
        for t in lb.get("tactics") or []:
            _add_heading(doc, t.get("name", "?"), 2)
            if t.get("target_links_per_month") is not None:
                _add_label_value(doc, "Target links / month", str(t["target_links_per_month"]))
            if t.get("rationale"):
                _add_body(doc, t["rationale"])
            if t.get("first_step"):
                _add_label_value(doc, "First step", t["first_step"])

    # ---------- Technical Priorities ----------
    tp = content.get("technical_priorities") or []
    if tp:
        _add_heading(doc, "Technical Priorities", 1)
        for t in tp:
            _add_heading(doc, t.get("title", "?"), 2)
            if t.get("rationale"):
                _add_body(doc, t["rationale"])
            if t.get("first_step"):
                _add_label_value(doc, "First step", t["first_step"])

    # ---------- 30/60/90 ----------
    ap = content.get("action_plan") or {}
    if ap:
        _add_heading(doc, "30 · 60 · 90 Day Action Plan", 1)
        for key, label in [("30_days", "First 30 days"), ("60_days", "Days 31-60"), ("90_days", "Days 61-90")]:
            items = ap.get(key) or []
            if not items:
                continue
            _add_heading(doc, label, 2, color=EMERALD)
            for it in items:
                _add_heading(doc, it.get("task", "?"), 3)
                _add_label_value(doc, "Owner", it.get("owner") or "—")
                _add_label_value(doc, "Deliverable", it.get("deliverable") or "—")

    # ---------- Success Metrics ----------
    metrics = content.get("success_metrics") or []
    if metrics:
        _add_heading(doc, "Success Metrics", 1)
        table = doc.add_table(rows=1, cols=4)
        for i, h in enumerate(["Metric", "Current", "90-day target", "180-day target"]):
            cell = table.rows[0].cells[i]
            cell.text = ""
            _set_cell_bg(cell, LIGHT_BG)
            _set_cell_border(cell)
            r = cell.paragraphs[0].add_run(h.upper())
            r.font.size = Pt(8)
            r.bold = True
            r.font.color.rgb = SLATE
        for m in metrics:
            row = table.add_row()
            for i, v in enumerate([m.get("metric"), m.get("current"), m.get("target_90d"), m.get("target_180d")]):
                cell = row.cells[i]
                cell.text = ""
                _set_cell_border(cell)
                r = cell.paragraphs[0].add_run(_fmt(v) if i > 0 else (v or "—"))
                r.font.size = Pt(10)
                r.font.color.rgb = EMERALD if i > 0 else CHARCOAL
                if i > 0:
                    r.bold = True

    # ---------- Closing ----------
    if content.get("closing_statement"):
        _add_heading(doc, "Start This Week", 1, color=EMERALD)
        _add_body(doc, content["closing_statement"], size=12)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ----------------------------- XLSX ------------------------------------------

def build_competitive_xlsx(content: Dict[str, Any], client_name: str = "") -> bytes:
    wb = Workbook()
    # Default sheet → Executive
    exec_ws = wb.active
    exec_ws.title = "Executive"

    HEADER_FONT = Font(name="Calibri", size=11, bold=True, color="0F172A")
    HEADER_FILL = PatternFill("solid", fgColor=LIGHT_BG)
    TITLE_FONT = Font(name="Calibri", size=20, bold=True, color="0F172A")
    SUBTITLE_FONT = Font(name="Calibri", size=11, color="475569", italic=True)
    BODY_FONT = Font(name="Calibri", size=10, color="1E293B")
    ACCENT_FONT = Font(name="Calibri", size=10, bold=True, color="B45309")
    EMERALD_FONT = Font(name="Calibri", size=10, bold=True, color="047857")
    BORDER = Border(*[Side(style="thin", color=DIVIDER)] * 4)
    WRAP = Alignment(wrap_text=True, vertical="top")

    def _autosize(ws, col, max_width=60):
        max_len = 0
        for cell in ws[col]:
            if cell.value:
                max_len = max(max_len, min(len(str(cell.value)), max_width))
        ws.column_dimensions[col].width = max(12, max_len + 2)

    # === Executive sheet ===
    exec_ws["A1"] = content.get("title") or f"Competitive Analysis · {client_name}"
    exec_ws["A1"].font = TITLE_FONT
    exec_ws.merge_cells("A1:E1")

    exec_ws["A2"] = content.get("subtitle") or ""
    exec_ws["A2"].font = SUBTITLE_FONT
    exec_ws["A2"].alignment = WRAP
    exec_ws.merge_cells("A2:E2")

    exec_ws["A3"] = f"Prepared for {content.get('prepared_for') or client_name}"
    exec_ws["A3"].font = Font(name="Calibri", size=9, color="94A3B8", italic=True)
    exec_ws.merge_cells("A3:E3")

    if content.get("executive_summary"):
        exec_ws["A5"] = "EXECUTIVE SUMMARY"
        exec_ws["A5"].font = ACCENT_FONT
        exec_ws["A6"] = content["executive_summary"]
        exec_ws["A6"].font = BODY_FONT
        exec_ws["A6"].alignment = WRAP
        exec_ws.merge_cells("A6:E6")
        exec_ws.row_dimensions[6].height = 80

    cp = content.get("current_position") or {}
    if cp:
        exec_ws["A8"] = "CURRENT POSITION"
        exec_ws["A8"].font = ACCENT_FONT
        headers = ["Authority", "Backlinks", "Ref. domains", "Organic KWs", "Organic traffic"]
        values = [cp.get("authority_score"), cp.get("backlinks"), cp.get("referring_domains"),
                  cp.get("organic_keywords"), cp.get("organic_traffic")]
        for i, h in enumerate(headers):
            c = exec_ws.cell(row=9, column=i + 1, value=h)
            c.font = HEADER_FONT
            c.fill = HEADER_FILL
            c.border = BORDER
            c.alignment = Alignment(horizontal="center")
            vc = exec_ws.cell(row=10, column=i + 1, value=values[i])
            vc.font = Font(name="Calibri", size=14, bold=True, color="0F172A")
            vc.border = BORDER
            vc.alignment = Alignment(horizontal="center")
        if cp.get("narrative"):
            exec_ws["A12"] = cp["narrative"]
            exec_ws["A12"].font = BODY_FONT
            exec_ws["A12"].alignment = WRAP
            exec_ws.merge_cells("A12:E12")
            exec_ws.row_dimensions[12].height = 50

    if content.get("closing_statement"):
        exec_ws["A14"] = "START THIS WEEK"
        exec_ws["A14"].font = EMERALD_FONT
        exec_ws["A15"] = content["closing_statement"]
        exec_ws["A15"].font = BODY_FONT
        exec_ws["A15"].alignment = WRAP
        exec_ws.merge_cells("A15:E15")
        exec_ws.row_dimensions[15].height = 50

    for col in "ABCDE":
        exec_ws.column_dimensions[col].width = 22

    # === Competitors sheet ===
    landscape = content.get("competitor_landscape") or []
    if landscape:
        ws = wb.create_sheet("Competitors")
        headers = ["Name", "Domain", "Tier", "Authority", "Organic traffic", "Strengths", "Weaknesses", "How to beat them"]
        for i, h in enumerate(headers):
            c = ws.cell(row=1, column=i + 1, value=h)
            c.font = HEADER_FONT
            c.fill = HEADER_FILL
            c.border = BORDER
        for r, comp in enumerate(landscape, start=2):
            row = [
                comp.get("name"),
                comp.get("domain"),
                (comp.get("tier") or "").replace("_", " "),
                comp.get("authority_score"),
                comp.get("organic_traffic"),
                " · ".join(comp.get("key_strengths") or []),
                " · ".join(comp.get("key_weaknesses") or []),
                comp.get("how_to_beat_them"),
            ]
            for i, v in enumerate(row):
                c = ws.cell(row=r, column=i + 1, value=v)
                c.font = BODY_FONT
                c.border = BORDER
                c.alignment = WRAP
            ws.row_dimensions[r].height = 80
        for col, w in zip("ABCDEFGH", [20, 28, 16, 12, 16, 50, 50, 50]):
            ws.column_dimensions[col].width = w

    # === Opportunities sheet ===
    opps = content.get("top_opportunities") or []
    if opps:
        ws = wb.create_sheet("Opportunities")
        headers = ["Rank", "Title", "Primary KW", "Volume", "Their position", "URL to outrank",
                   "Format", "Effort", "Impact", "Why winnable", "First step"]
        for i, h in enumerate(headers):
            c = ws.cell(row=1, column=i + 1, value=h)
            c.font = HEADER_FONT
            c.fill = HEADER_FILL
            c.border = BORDER
        for r, o in enumerate(opps, start=2):
            row = [
                o.get("rank"),
                o.get("title"),
                o.get("primary_keyword"),
                o.get("search_volume"),
                o.get("competitor_position"),
                o.get("competitor_url_to_outrank"),
                o.get("recommended_format"),
                o.get("effort"),
                o.get("expected_impact"),
                o.get("why_winnable"),
                o.get("first_step"),
            ]
            for i, v in enumerate(row):
                c = ws.cell(row=r, column=i + 1, value=v)
                c.font = BODY_FONT
                c.border = BORDER
                c.alignment = WRAP
            ws.row_dimensions[r].height = 60
        widths = [6, 36, 24, 10, 12, 36, 22, 10, 10, 40, 40]
        for col, w in zip("ABCDEFGHIJK", widths):
            ws.column_dimensions[col].width = w

    # === Action plan sheet ===
    ap = content.get("action_plan") or {}
    if ap:
        ws = wb.create_sheet("Action Plan")
        for i, h in enumerate(["Bucket", "Task", "Owner", "Deliverable"]):
            c = ws.cell(row=1, column=i + 1, value=h)
            c.font = HEADER_FONT
            c.fill = HEADER_FILL
            c.border = BORDER
        r = 2
        for key, label in [("30_days", "First 30 days"), ("60_days", "Days 31-60"), ("90_days", "Days 61-90")]:
            for it in ap.get(key) or []:
                for i, v in enumerate([label, it.get("task"), it.get("owner"), it.get("deliverable")]):
                    c = ws.cell(row=r, column=i + 1, value=v)
                    c.font = BODY_FONT if i > 0 else ACCENT_FONT
                    c.border = BORDER
                    c.alignment = WRAP
                ws.row_dimensions[r].height = 40
                r += 1
        for col, w in zip("ABCD", [16, 50, 24, 50]):
            ws.column_dimensions[col].width = w

    # === Success metrics sheet ===
    metrics = content.get("success_metrics") or []
    if metrics:
        ws = wb.create_sheet("Success Metrics")
        for i, h in enumerate(["Metric", "Current", "90-day target", "180-day target"]):
            c = ws.cell(row=1, column=i + 1, value=h)
            c.font = HEADER_FONT
            c.fill = HEADER_FILL
            c.border = BORDER
        for r, m in enumerate(metrics, start=2):
            for i, v in enumerate([m.get("metric"), m.get("current"), m.get("target_90d"), m.get("target_180d")]):
                c = ws.cell(row=r, column=i + 1, value=v)
                c.font = EMERALD_FONT if i > 0 else BODY_FONT
                c.border = BORDER
                c.alignment = Alignment(horizontal="right" if i > 0 else "left")
        for col, w in zip("ABCD", [32, 16, 18, 18]):
            ws.column_dimensions[col].width = w

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()
